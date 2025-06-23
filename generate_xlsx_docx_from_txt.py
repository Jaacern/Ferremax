import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font
from docx import Document
import os
from openpyxl.cell.cell import MergedCell

# Paths
BASE = os.path.dirname(os.path.abspath(__file__))
TXT_CASOS = os.path.join(BASE, 'CASOS_PRUEBA_INTEGRACION.txt')
TXT_DEFECTOS = os.path.join(BASE, 'REGISTRO_DEFECTOS.txt')
TXT_PLAN = os.path.join(BASE, 'PLAN_PRUEBAS.txt')
PLANTILLA_CASOS = os.path.join(BASE, '3.1.4 Plantilla Casos de prueba Integracion.xlsx')
PLANTILLA_DEFECTOS = os.path.join(BASE, '3.3.4 Planilla_Registro_de_Defectos_ejemplo.xlsx')
PLANTILLA_PLAN = os.path.join(BASE, '3.2.4 Plantilla_Plan_de_Pruebas.docx')
OUT_CASOS = os.path.join(BASE, '3.1.4 Casos_de_prueba_Integracion_COMPLETO.xlsx')
OUT_DEFECTOS = os.path.join(BASE, '3.3.4 Registro_de_Defectos_COMPLETO.xlsx')
OUT_PLAN = os.path.join(BASE, '3.2.4 Plan_de_Pruebas_COMPLETO.docx')

def set_cell_value(ws, row, col, value, bold=False):
    # Busca la primera celda no combinada desde la fila dada
    max_row = ws.max_row
    while row <= max_row:
        cell = ws.cell(row=row, column=col)
        if not isinstance(cell, MergedCell):
            cell.value = value
            if bold:
                cell.font = Font(bold=True)
            return
        row += 1

def write_table(ws, header, data):
    # Escribir encabezado
    for col, h in enumerate(header, 1):
        set_cell_value(ws, 1, col, h, bold=True)
    # Escribir datos
    for i, row in enumerate(data, 2):
        for j, val in enumerate(row, 1):
            set_cell_value(ws, i, j, val)

# 1. CASOS DE PRUEBA INTEGRACION
if os.path.exists(PLANTILLA_CASOS):
    wb = openpyxl.load_workbook(PLANTILLA_CASOS)
    ws = wb.active if wb.active else wb.worksheets[0]
    # Leer datos del txt
    with open(TXT_CASOS, encoding='utf-8') as f:
        lines = [l.strip() for l in f if l.strip() and not l.startswith('=') and not l.startswith('CASOS DE PRUEBA')]
    # Buscar encabezado y datos
    header_idx = next(i for i, l in enumerate(lines) if l.startswith('ID'))
    header = lines[header_idx].split('\t')
    data = [l.split('\t') for l in lines[header_idx+1:] if l.count('\t') >= len(header)-1]
    # Limpiar hoja (excepto encabezado)
    max_row = ws.max_row
    if max_row > 1:
        for row in ws.iter_rows(min_row=2, max_row=max_row):
            for cell in row:
                if not isinstance(cell, MergedCell):
                    cell.value = None
    write_table(ws, header, data)
    wb.save(OUT_CASOS)
    print(f"Archivo generado: {OUT_CASOS}")
else:
    print("No se encontró la plantilla de casos de prueba.")

# 2. REGISTRO DE DEFECTOS
if os.path.exists(PLANTILLA_DEFECTOS):
    wb = openpyxl.load_workbook(PLANTILLA_DEFECTOS)
    ws = wb.active if wb.active else wb.worksheets[0]
    with open(TXT_DEFECTOS, encoding='utf-8') as f:
        lines = [l.strip() for l in f if l.strip() and not l.startswith('=') and not l.startswith('REGISTRO DE DEFECTOS')]
    header_idx = next(i for i, l in enumerate(lines) if l.startswith('ID'))
    header = lines[header_idx].split('\t')
    data = [l.split('\t') for l in lines[header_idx+1:] if l.count('\t') >= len(header)-1]
    max_row = ws.max_row
    if max_row > 1:
        for row in ws.iter_rows(min_row=2, max_row=max_row):
            for cell in row:
                if not isinstance(cell, MergedCell):
                    cell.value = None
    write_table(ws, header, data)
    wb.save(OUT_DEFECTOS)
    print(f"Archivo generado: {OUT_DEFECTOS}")
else:
    print("No se encontró la plantilla de registro de defectos.")

# 3. PLAN DE PRUEBAS
if os.path.exists(PLANTILLA_PLAN):
    doc = Document(PLANTILLA_PLAN)
    # Leer el texto plano
    with open(TXT_PLAN, encoding='utf-8') as f:
        plan_text = f.read()
    # Limpiar el documento (eliminar todos los párrafos)
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p.clear()
        p._element.getparent().remove(p._element)
    # Agregar el texto del plan
    for line in plan_text.splitlines():
        doc.add_paragraph(line)
    doc.save(OUT_PLAN)
    print(f"Archivo generado: {OUT_PLAN}")
else:
    print("No se encontró la plantilla de plan de pruebas.") 